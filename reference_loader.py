import io
from pathlib import Path

from docx import Document


REFERENCE_DIR = Path(__file__).resolve().parent / "references"
DEFAULT_INSTRUCTIONS = REFERENCE_DIR / "Odigies_v5.docx"
DEFAULT_STYLE = REFERENCE_DIR / "Elena_style_guide_v2.docx"
ROOT_INSTRUCTIONS = Path(__file__).resolve().parent / "Odigies_v5.docx"
ROOT_STYLE = Path(__file__).resolve().parent / "Elena_style_guide_v2.docx"


def docx_text(source) -> str:
    """Extract paragraphs and tables from a DOCX path or uploaded bytes."""
    if isinstance(source, (bytes, bytearray)):
        source = io.BytesIO(source)
    document = Document(source)
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            line = " | ".join(cell.text.strip() for cell in row.cells)
            if line.strip(" |"):
                blocks.append(line)
    return "\n".join(blocks)


def load_default_references() -> tuple[str, str]:
    # Accept both repository layouts: a dedicated references/ folder or the
    # two DOCX files beside app.py.  This makes GitHub web uploads simpler.
    instructions = DEFAULT_INSTRUCTIONS if DEFAULT_INSTRUCTIONS.exists() else ROOT_INSTRUCTIONS
    style = DEFAULT_STYLE if DEFAULT_STYLE.exists() else ROOT_STYLE
    if not instructions.exists() or not style.exists():
        raise FileNotFoundError("Λείπουν οι ενσωματωμένες οδηγίες v4 ή το πρότυπο ύφους.")
    return docx_text(instructions), docx_text(style)
