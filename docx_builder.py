import re
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from prompts import fmt
from astrology import movement_text

def _shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

# Αναγνωρίζει **έντονα**, *πλάγια* και ***και τα δύο***· δεν πειράζει απλό
# κείμενο χωρίς αστερίσκους. Χρησιμοποιείται αντί για αφαίρεση των αστερίσκων,
# ώστε η έμφαση του μοντέλου να φτάνει πραγματικά στο Word.
_INLINE_MD = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")
_NUMBERED_LIST = re.compile(r"^(\d+)\.\s+(.+)$")

def _add_formatted_runs(paragraph, text):
    for chunk in _INLINE_MD.split(text):
        if not chunk:
            continue
        if chunk.startswith('***') and chunk.endswith('***') and len(chunk) > 6:
            r = paragraph.add_run(chunk[3:-3]); r.bold = True; r.italic = True
        elif chunk.startswith('**') and chunk.endswith('**') and len(chunk) > 4:
            r = paragraph.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith('*') and chunk.endswith('*') and len(chunk) > 2:
            r = paragraph.add_run(chunk[1:-1]); r.italic = True
        else:
            paragraph.add_run(chunk)

def _add_multiline(doc, text):
    """Προσθέτει το `text` ως ένα Word paragraph ανά γραμμή.

    Το python-docx (και το OOXML γενικότερα) ΔΕΝ μετατρέπει ένα literal '\\n'
    μέσα σε run σε ορατή αλλαγή γραμμής -- το Word δείχνει ένα συνεχόμενο
    μπλοκ κειμένου. Το build_audit_docx καλούσε d.add_paragraph(prompt) μία
    φορά με ολόκληρη την πολυσέλιδη εντολή (όλοι οι 12 Οίκοι, οι πλήρεις
    οδηγίες v4, το πρότυπο ύφους) ως ένα string, οπότε έβγαινε ένα άμορφο
    μπλοκ ~88.000 χαρακτήρων χωρίς ορατά όρια ενοτήτων. Εδώ σπάει σε πραγματικές
    γραμμές πρώτα, ώστε κάθε γραμμή -- ακόμη και οι κενές, που γίνονται κενά
    paragraphs -- να είναι δικό της paragraph, όπως ακριβώς η υπόλοιπη δομή
    αυτού του εγγράφου.
    """
    for line in text.splitlines():
        doc.add_paragraph(line)

def build_audit_docx(chart, personal, prompt):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    styles=d.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10.5)
    for s,size,color in [('Title',26,'1D3A34'),('Heading 1',18,'1D3A34'),('Heading 2',14,'5B7F6A')]:
        styles[s].font.name='Aptos Display'; styles[s].font.size=Pt(size); styles[s].font.color.rgb=RGBColor.from_string(color)
    title=d.add_paragraph(style='Title'); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.add_run('AstroCheck — Δελτίο ελέγχου')
    # Το διορθωμένο όνομα (personal["Όνομα"]) είναι η πηγή αλήθειας, ίδια
    # λογική με το build_master_prompt -- ώστε ο τίτλος αυτού του Δελτίου
    # Ελέγχου να συμφωνεί πάντα με ό,τι βλέπει το μοντέλο μέσα στην εντολή.
    display_name = (personal.get("Όνομα") or "").strip() or chart.name
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(display_name).bold=True
    d.add_heading('Βασικά δεδομένα',1)
    table=d.add_table(rows=0,cols=2); table.style='Table Grid'
    for a,b in [('Ημερομηνία',chart.date),('Ώρα',chart.time),('Τόπος',chart.place),('Σύστημα Οίκων',chart.house_system)]:
        cells=table.add_row().cells; cells[0].text=a; cells[1].text=b; _shade(cells[0],'E5EFE8')
    d.add_heading('Πλανήτες και σημεία',1)
    t=d.add_table(rows=1,cols=4); t.style='Table Grid'
    for i,h in enumerate(['Σημείο','Θέση','Οίκος','Κίνηση']): t.rows[0].cells[i].text=h; _shade(t.rows[0].cells[i],'1D3A34'); t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
    for pnt in chart.points:
        c=t.add_row().cells; c[0].text=pnt.name; c[1].text=fmt(pnt); c[2].text=str(pnt.house or '—'); c[3].text=movement_text(pnt)
    d.add_heading('Υποχρεωτικός έλεγχος τετραγώνων και αντιθέσεων',1)
    hard=[a for a in chart.aspects if a.aspect in ('Τετράγωνο','Αντίθεση')]
    t=d.add_table(rows=1,cols=4); t.style='Table Grid'
    for i,h in enumerate(['Ζεύγος','Όψη','Orb','Βαρύτητα']): t.rows[0].cells[i].text=h; _shade(t.rows[0].cells[i],'1D3A34'); t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
    for a in hard:
        c=t.add_row().cells; c[0].text=f'{a.first}–{a.second}'; c[1].text=a.aspect; c[2].text=a.orb_text; c[3].text=a.weight
    d.add_page_break(); d.add_heading('Πλήρης εντολή για δημιουργία ανάλυσης',1)
    _add_multiline(d, prompt)
    bio=BytesIO(); d.save(bio); return bio.getvalue()

_HOUSE_HEADING_RE = re.compile(r'^\d{1,2}ος\s+Ο[ιί]κος', re.IGNORECASE)
_APPENDIX_HEADING_RE = re.compile(r'^Παράρτημα\s+επιβεβαιωμένων\s+όψεων', re.IGNORECASE)
_SUMMARY_START_RE = re.compile(r'^Βασική\s+δύναμη\s*[:\-]', re.IGNORECASE)
_SUMMARY_LABELS = ('Βασική δύναμη', 'Βασική πρόκληση', 'Κυβερνήτης', 'Τελικό συμπέρασμα')


def _add_page_number_field(paragraph):
    """PAGE / NUMPAGES στο footer -- το Odigies (§18) ζητά αριθμημένες
    σελίδες· το build_analysis_docx δεν το έκανε καθόλου πριν."""
    def field(instr):
        run = paragraph.add_run()
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
        run._r.append(f1); run._r.append(it); run._r.append(f2)
    field('PAGE')
    paragraph.add_run(' / ')
    field('NUMPAGES')


def _clean_summary_label(line):
    """Αφαιρεί bullet/markdown από μια γραμμή πλαισίου σύνοψης, κρατώντας
    'Ετικέτα: τιμή' έτοιμο για bold-label εμφάνιση."""
    line = line.lstrip('-•* ').strip()
    line = line.replace('**', '')
    return line


def _add_summary_box(doc, lines):
    """Πλαίσιο σύνοψης Οίκου ως πίνακας ενός κελιού με cantSplit -- έτσι
    μένει εγγυημένα ολόκληρο στην ίδια σελίδα (Odigies §11/§18), αντί για
    απλές παραγράφους που το Word μπορεί να σπάσει ανάμεσα σε δύο σελίδες."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    row = table.rows[0]
    row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    cell = row.cells[0]
    _shade(cell, 'EEF2F6')
    cell.paragraphs[0].text = ''
    first = True
    for raw in lines:
        text = _clean_summary_label(raw)
        if not text:
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        label_matched = False
        for label in _SUMMARY_LABELS:
            if text.lower().startswith(label.lower()):
                rest = text[len(label):].lstrip(': ').strip()
                r = p.add_run(label + ': '); r.bold = True
                _add_formatted_runs(p, rest)
                label_matched = True
                break
        if not label_matched:
            _add_formatted_runs(p, text)
    doc.add_paragraph()  # μικρό κενό μετά το πλαίσιο, πριν τον επόμενο Οίκο


def build_analysis_docx(title_name, analysis):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
    d.styles['Normal'].font.name='Aptos'; d.styles['Normal'].font.size=Pt(10.5)
    for s,size,color in [('Title',26,'1D3A34'),('Heading 1',18,'1D3A34'),('Heading 2',14,'5B7F6A'),('Heading 3',12,'5B7F6A')]:
        d.styles[s].font.name='Aptos Display'; d.styles[s].font.size=Pt(size); d.styles[s].font.color.rgb=RGBColor.from_string(color)
    p=d.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Πλήρης Αστρολογική Ανάλυση')
    s=d.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER; s.add_run(title_name).bold=True

    footer_p = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_p)

    lines = analysis.splitlines()
    i = 0
    first_house_seen = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        core = line
        for prefix in ('### ', '## ', '# '):
            if core.startswith(prefix):
                core = core[len(prefix):]
                break
        core_plain = core.replace('**', '').lstrip('-•* ').strip()

        # Το παλιό υποχρεωτικό page break πριν από κάθε Οίκο δημιουργούσε
        # κενές ή σχεδόν κενές σελίδες όταν το πλαίσιο σύνοψης του προηγούμενου
        # Οίκου μεταφερόταν ολόκληρο στην επόμενη σελίδα. Οι οδηγίες ζητούν νέα
        # σελίδα «κατά προτίμηση», όχι εις βάρος της σελιδοποίησης. Αφήνουμε
        # πλέον φυσική ροή μεταξύ Οίκων και κρατάμε μόνο το Παράρτημα σε νέα
        # σελίδα.
        if _HOUSE_HEADING_RE.match(core_plain):
            first_house_seen = True
        elif _APPENDIX_HEADING_RE.match(core_plain):
            # Φυσική ροή και εδώ: ένα ρητό page break μπορεί να δημιουργήσει
            # ολόκληρη κενή σελίδα όταν το προηγούμενο περιεχόμενο έχει ήδη
            # γεμίσει ακριβώς την τρέχουσα σελίδα.
            pass

        # Πλαίσιο σύνοψης: μόλις εντοπιστεί η πρώτη ετικέτα, μαζεύουμε τις
        # επόμενες γραμμές μέχρι την πρώτη κενή γραμμή και τις ρίχνουμε σε
        # έναν πίνακα-πλαίσιο αντί για απλές παραγράφους.
        if _SUMMARY_START_RE.match(core_plain):
            box_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip():
                box_lines.append(lines[j].strip())
                j += 1
            _add_summary_box(d, box_lines)
            i = j
            continue

        if not line:
            d.add_paragraph(); i += 1; continue
        if line.startswith('### '): _add_formatted_runs(d.add_heading('',3), line[4:])
        elif line.startswith('## '): _add_formatted_runs(d.add_heading('',2), line[3:])
        elif line.startswith('# '): _add_formatted_runs(d.add_heading('',1), line[2:])
        elif line.startswith(('- ','• ')): _add_formatted_runs(d.add_paragraph(style='List Bullet'), line[2:])
        else:
            m = _NUMBERED_LIST.match(line)
            if m: _add_formatted_runs(d.add_paragraph(style='List Number'), m.group(2))
            else: _add_formatted_runs(d.add_paragraph(), line)
        i += 1
    bio=BytesIO(); d.save(bio); return bio.getvalue()
